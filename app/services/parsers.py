from __future__ import annotations

import io
import mimetypes
import re
from dataclasses import dataclass

import docx
from PIL import Image

try:
    import pymupdf as fitz  # PyMuPDF new import
except Exception:  # pragma: no cover
    fitz = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    import pytesseract
except Exception:  # pragma: no cover
    pytesseract = None


class UnsupportedFileTypeError(Exception):
    pass


@dataclass
class ParseResult:
    text: str
    mime_type: str
    parser_meta: dict


MIN_REASONABLE_TEXT = 80
_WHITESPACE_RE = re.compile(r"\s+")
_FONT_BOLD_RE = re.compile(r"bold|black|heavy|demi", re.I)
_IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp", ".bmp")


def _guess_mime_type(filename: str) -> str:
    mime_type, _ = mimetypes.guess_type(filename)
    return mime_type or "application/octet-stream"


def _normalize_text(text: str) -> str:
    lines = [line.rstrip() for line in text.splitlines()]
    collapsed = []
    previous_blank = False
    for line in lines:
        clean = " ".join(line.split())
        if not clean:
            if not previous_blank:
                collapsed.append("")
            previous_blank = True
            continue
        previous_blank = False
        collapsed.append(clean)
    return "\n".join(collapsed).strip()


def _group_lines_by_column(layout_lines: list[dict], page_width: float) -> list[dict]:
    if not layout_lines:
        return []

    candidate_xs = sorted(
        {
            round(float(item.get("x0") or 0.0), 1)
            for item in layout_lines
            if float(item.get("font_size") or 0.0) >= 13.0 or bool(item.get("is_bold"))
        }
    )
    split_x = page_width * 0.5 if page_width > 0 else 220.0
    if len(candidate_xs) >= 2:
        gaps = [
            (candidate_xs[i + 1] - candidate_xs[i], candidate_xs[i], candidate_xs[i + 1])
            for i in range(len(candidate_xs) - 1)
        ]
        best_gap, left_x, right_x = max(gaps, key=lambda item: item[0])
        if best_gap >= 60:
            split_x = (left_x + right_x) / 2.0

    left = [item for item in layout_lines if float(item.get("x0") or 0.0) < split_x]
    right = [item for item in layout_lines if float(item.get("x0") or 0.0) >= split_x]

    if not left or not right:
        return sorted(layout_lines, key=lambda item: (round(float(item["y0"]), 1), round(float(item["x0"]), 1)))

    return sorted(left, key=lambda item: (round(float(item["y0"]), 1), round(float(item["x0"]), 1))) + sorted(
        right,
        key=lambda item: (round(float(item["y0"]), 1), round(float(item["x0"]), 1)),
    )


def _reconstruct_text_from_layout(layout_lines: list[dict], page_dimensions: dict[int, tuple[float, float]] | None = None) -> str:
    if not layout_lines:
        return ""

    page_dimensions = page_dimensions or {}
    pages: dict[int, list[dict]] = {}
    for line in layout_lines:
        pages.setdefault(int(line["page"]), []).append(line)

    output: list[str] = []
    for page in sorted(pages.keys()):
        page_lines = pages[page]
        page_width = float(page_dimensions.get(page, (0.0, 0.0))[0] or 0.0)
        ordered_lines = _group_lines_by_column(page_lines, page_width)

        prev_y: float | None = None
        prev_x: float | None = None
        for line in ordered_lines:
            y0 = float(line["y0"])
            x0 = float(line["x0"])
            if prev_y is not None:
                gap = y0 - prev_y
                if gap > max(14.0, float(line["font_size"]) * 1.2):
                    output.append("")
                elif prev_x is not None and abs(x0 - prev_x) > max(140.0, page_width * 0.18 if page_width else 140.0):
                    output.append("")
            output.append(line["text"])
            prev_y = y0
            prev_x = x0
        output.append("")

    return _normalize_text("\n".join(output))


def _extract_pdf_text_pymupdf(file_bytes: bytes) -> tuple[str, dict]:
    if fitz is None:
        raise RuntimeError("PyMuPDF is not installed")

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        layout_lines: list[dict] = []
        first_page_layout_lines: list[dict] = []
        page_dimensions: dict[int, tuple[float, float]] = {}

        for page_index, page in enumerate(doc):
            page_no = page_index + 1
            page_dimensions[page_no] = (float(page.rect.width), float(page.rect.height))

            page_dict = page.get_text("dict", sort=True)
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue

                for line in block.get("lines", []):
                    spans = [span for span in line.get("spans", []) if span.get("text", "").strip()]
                    if not spans:
                        continue

                    spans = sorted(spans, key=lambda item: item["bbox"][0])
                    text = "".join(span.get("text", "") for span in spans)
                    text = _WHITESPACE_RE.sub(" ", text).strip()
                    if not text:
                        continue

                    x0 = min(float(span["bbox"][0]) for span in spans)
                    y0 = min(float(span["bbox"][1]) for span in spans)
                    x1 = max(float(span["bbox"][2]) for span in spans)
                    y1 = max(float(span["bbox"][3]) for span in spans)
                    font_size = max(float(span.get("size", 0.0)) for span in spans)
                    font_names = [str(span.get("font", "")) for span in spans]
                    is_bold = any(_FONT_BOLD_RE.search(name or "") for name in font_names)

                    line_item = {
                        "page": page_no,
                        "text": text,
                        "x0": round(x0, 2),
                        "y0": round(y0, 2),
                        "x1": round(x1, 2),
                        "y1": round(y1, 2),
                        "font_size": round(font_size, 2),
                        "is_bold": bool(is_bold),
                    }
                    layout_lines.append(line_item)
                    if page_index == 0:
                        first_page_layout_lines.append(line_item)

        text = _reconstruct_text_from_layout(layout_lines, page_dimensions=page_dimensions)
        first_page_width, first_page_height = page_dimensions.get(1, (0.0, 0.0))
        meta = {
            "parser": "pymupdf_layout",
            "used_ocr": False,
            "page_count": len(doc),
            "low_text_warning": len(text) < MIN_REASONABLE_TEXT,
            "first_page_width": first_page_width,
            "first_page_height": first_page_height,
            "first_page_layout_lines": first_page_layout_lines[:250],
            "all_page_layout_lines": layout_lines[:3000],
            "page_dimensions": {
                str(page): {"width": width, "height": height}
                for page, (width, height) in page_dimensions.items()
            },
        }
        return text, meta
    finally:
        doc.close()


def _extract_pdf_text_pypdf(file_bytes: bytes) -> tuple[str, dict]:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed")

    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    text = _normalize_text("\n\n".join(pages))
    meta = {
        "parser": "pypdf",
        "used_ocr": False,
        "page_count": len(pages),
        "low_text_warning": len(text) < MIN_REASONABLE_TEXT,
        "first_page_layout_lines": [],
        "all_page_layout_lines": [],
        "first_page_width": 0.0,
        "first_page_height": 0.0,
        "page_dimensions": {},
    }
    return text, meta


def _ocr_image_pytesseract(image: Image.Image, *, lang: str = "vie+eng") -> str:
    if pytesseract is None:
        raise RuntimeError("pytesseract is not installed")
    rgb = image.convert("RGB")
    return _normalize_text(pytesseract.image_to_string(rgb, lang=lang, config="--oem 1 --psm 6"))


def _extract_pdf_text(file_bytes: bytes) -> tuple[str, dict]:
    primary_error = None
    try:
        if fitz is not None:
            text, meta = _extract_pdf_text_pymupdf(file_bytes)
            if text.strip():
                return text, meta
    except Exception as exc:  # pragma: no cover
        primary_error = str(exc)

    text, meta = _extract_pdf_text_pypdf(file_bytes)
    if primary_error:
        meta["fallback_reason"] = primary_error
    return text, meta


def _extract_docx_text(file_bytes: bytes) -> tuple[str, dict]:
    document = docx.Document(io.BytesIO(file_bytes))

    paragraphs = []
    layout_lines = []

    for index, paragraph in enumerate(document.paragraphs):
        text = " ".join(paragraph.text.split()).strip()
        if not text:
            continue

        font_sizes = []
        is_bold = False
        for run in paragraph.runs:
            if run.font and run.font.size:
                try:
                    font_sizes.append(float(run.font.size.pt))
                except Exception:
                    pass
            if bool(run.bold):
                is_bold = True

        font_size = max(font_sizes) if font_sizes else 12.0
        paragraphs.append(text)
        if index < 180:
            layout_lines.append(
                {
                    "page": 1,
                    "text": text,
                    "x0": 0.0,
                    "y0": float(index * 18),
                    "x1": 0.0,
                    "y1": float(index * 18 + 16),
                    "font_size": round(font_size, 2),
                    "is_bold": is_bold,
                }
            )

    tables: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables.append(" | ".join(cells))

    text = _normalize_text("\n".join(paragraphs + tables))
    meta = {
        "parser": "python-docx",
        "used_ocr": False,
        "paragraph_count": len(paragraphs),
        "table_row_count": len(tables),
        "low_text_warning": len(text) < MIN_REASONABLE_TEXT,
        "first_page_layout_lines": layout_lines[:250],
        "all_page_layout_lines": layout_lines[:800],
        "first_page_width": 0.0,
        "first_page_height": 0.0,
        "page_dimensions": {"1": {"width": 0.0, "height": 0.0}},
    }
    return text, meta


def _extract_txt_text(file_bytes: bytes) -> tuple[str, dict]:
    text = _normalize_text(file_bytes.decode("utf-8", errors="ignore"))
    meta = {
        "parser": "plain-text",
        "used_ocr": False,
        "low_text_warning": len(text) < MIN_REASONABLE_TEXT,
        "first_page_layout_lines": [],
        "all_page_layout_lines": [],
        "first_page_width": 0.0,
        "first_page_height": 0.0,
        "page_dimensions": {},
    }
    return text, meta


def _extract_image_text(file_bytes: bytes) -> tuple[str, dict]:
    image = Image.open(io.BytesIO(file_bytes))
    text = _ocr_image_pytesseract(image)
    width, height = image.size
    meta = {
        "parser": "image_ocr_tesseract",
        "used_ocr": True,
        "low_text_warning": len(text) < MIN_REASONABLE_TEXT,
        "first_page_layout_lines": [],
        "all_page_layout_lines": [],
        "first_page_width": float(width),
        "first_page_height": float(height),
        "page_dimensions": {"1": {"width": float(width), "height": float(height)}},
    }
    return text, meta


def extract_text(file_bytes: bytes, filename: str) -> ParseResult:
    name = filename.lower()
    mime_type = _guess_mime_type(filename)

    if name.endswith(".pdf"):
        text, meta = _extract_pdf_text(file_bytes)
    elif name.endswith(".docx"):
        text, meta = _extract_docx_text(file_bytes)
    elif name.endswith(".txt"):
        text, meta = _extract_txt_text(file_bytes)
    elif name.endswith(_IMAGE_EXTENSIONS):
        text, meta = _extract_image_text(file_bytes)
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {filename}")

    return ParseResult(text=text, mime_type=mime_type, parser_meta=meta)